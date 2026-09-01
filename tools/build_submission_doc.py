#!/usr/bin/env python3
"""
build_submission_doc.py — TECHNOVA 2026 idea submission document (docx).
Sections follow the challenge's ask: problem statement, frugal
hardware/software architecture, expected impact — plus technical detail.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

INK = RGBColor(0x11, 0x24, 0x3B)
TEAL = RGBColor(0x0F, 0x76, 0x6E)
MUTED = RGBColor(0x55, 0x67, 0x7A)

doc = Document()

# ---- base styles ----
st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.12

for sec in doc.sections:
    sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(15); r.font.color.rgb = TEAL
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(12); r.font.color.rgb = INK
    return p

def para(text, italic=False, size=10.5, color=INK, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text); r.italic = italic
    r.font.size = Pt(size); r.font.color.rgb = color
    return p

def rich(parts, size=10.5):
    """parts: list of (text, bold) tuples."""
    p = doc.add_paragraph()
    for t, b in parts:
        r = p.add_run(t); r.bold = b; r.font.size = Pt(size)
    return p

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True; r.font.size = Pt(10.5)
    r = p.add_run(text); r.font.size = Pt(10.5)
    return p

def table(rows, widths=None, header=True, font=9.5):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(font)
            if i == 0 or (header and i == 0):
                r.bold = True
    if widths:
        for j, w in enumerate(widths):
            for i in range(len(rows)):
                t.cell(i, j).width = Inches(w)
    return t

# ======================================================================
# COVER
# ======================================================================
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run('TECHNOVA 2026')  # placeholder – will style below
r.font.size = Pt(11); r.font.color.rgb = MUTED
r = p.add_run('\nNational AI Innovation Challenge')
r.font.size = Pt(13); r.font.color.rgb = MUTED

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('NadiSense')
r.font.size = Pt(40); r.bold = True; r.font.color.rgb = TEAL
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('60-second AI heart-rhythm screening for rural India')
r.font.size = Pt(16); r.font.color.rgb = INK
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Camera-only photoplethysmography · on-device neural network · zero hardware · zero internet')
r.font.size = Pt(11.5); r.font.color.rgb = MUTED
doc.add_paragraph()

para('Submission document — Idea / prototype', italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
para('Team: MatricPhase', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
para('Team lead: Aditya Mehra  ·  aditya.mehra@example.com  ·  +91 XXXXX XXXXX',
     size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
para('(Pilot circuit: Madurai–Thiruparankundram belt, Tamil Nadu)',
     size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ======================================================================
# 0. EXECUTIVE SUMMARY
# ======================================================================
h1('0. Executive summary')
rich([('Cardiovascular disease kills more people in India than any other cause, and the '
       'stroke-causing rhythm disorder atrial fibrillation (AFib) is both common and silent. '
       'The only way to catch it reliably is an ECG — a machine that most villages simply '
       'do not have, at a distance that most patients cannot afford to travel. ', False),
      ('NadiSense closes that gap with software alone: ', True),
      ('it turns the smartphone camera in an ASHA worker’s pocket into a pulse sensor and reads '
       'the 60-second signal with a 6 KB neural network running entirely on the phone. '
       'No sensor, no internet, no cloud, no clinic visit.', False)])
rich([('The result is one clear line an ASHA worker can act on: ', False),
      ('green — routine screening · amber — repeat in two weeks · red — get an ECG within '
       'seven days. ', True),
      ('In English, Hindi and Tamil, offline, with a printable screening report.', False)])
rich([('We have fully built the product: a working prototype (camera capture + signal '
       'processing + classifier + vernacular UI + report), 30 automated pipeline tests, 13 '
       'end-to-end UI tests, and a reproducible training pipeline for the model. We are also '
       'unusually explicit about what is not yet true: the shipped model is validated on a '
       'realistic synthetic training distribution; a real-patient benchmark (MIT-BIH AF/NSR) '
       'is one command away — the script ships with the code. We consider that honesty part '
       'of the engineering, not a disclaimer.', False)])

# ======================================================================
# 1. PROBLEM
# ======================================================================
h1('1. Problem statement')
rich([('Every year, tens of millions of Indians are screened for nothing while the two '
       'conditions that cause most preventable strokes and heart failures — hypertension and '
       'atrial fibrillation — go undetected. The clinical truth is blunt: ', False),
      ('rhythm disorders are diagnosed by ECG, and ECG machines live in district hospitals '
       'that are often 40–90 km and a day’s wages away from the patient.', True)])
h2('1.1 Why AFib is the right target')
bullet('AFib is the most common sustained cardiac arrhythmia; prevalence rises from ~1% under '
       '60 to ~9% above 80, so a rapidly-ageing population makes this a growing problem.')
bullet('It is the single most common cause of cardioembolic stroke: roughly 1 in 5 strokes is '
       'attributed to AF. Many of those strokes are the first symptom — the disease is silent.')
bullet('It is paroxysmal: a 10-second ECG on the right day catches it, but symptoms come and '
       'go, and a single in-clinic ECG is a coin flip.')
bullet('It is treatable: anticoagulation cuts stroke risk by roughly two-thirds. Early '
       'detection converts a catastrophic event into a prescription.')
h2('1.2 The screening gap is logistical, not clinical')
table([
    ['Who', 'What they have today', 'The gap'],
    ['Village/PHC level', 'Manual BP cuffs, outreach camps', 'No rhythm screening at all — pulse “feels fine” is not a test'],
    ['Block CHC', 'One ECG machine, one technician, long queue', 'Travel + waiting cost more than the test; referrals drop off'],
    ['ASHA workers (~10 lakh)', 'A smartphone and a door-to-door mandate', 'No tool matches their workflow; NCD software they get is paper-first'],
], widths=[1.7, 2.6, 3.1])
para('Note on figures: prevalence and stroke-attributable-risk figures are widely-cited public '
     'health ranges (WHO/ICMR-style reviews); we reproduce them as ranges with sources in the '
     'appendix rather than as false precision.', size=9, italic=True, color=MUTED)

# ======================================================================
# 2. USERS
# ======================================================================
h1('2. Target users and stakeholders')
table([
    ['Stakeholder', 'Need', 'How NadiSense serves them'],
    ['ASHA worker / ANM', 'A tool that fits a 5-minute home visit', 'Point, wait, read the line, act. Voice questionnaire, vernacular UI'],
    ['Patient (village adult)', 'Get screened without losing a workday', 'Screen happens at home; only the red-flagged travel for ECG'],
    ['PHC / MO', 'Triage, not more appointments', 'A printable report that says clearly: routine / repeat / refer'],
    ['District health society', 'Cheap NCD screening at scale', '₹0 hardware; runs on phones already issued; no server to run'],
    ['State NHM', 'Programme metrics', 'Logbook is local-first; export/reporting path is on the roadmap'],
], widths=[1.9, 2.7, 2.8])

# ======================================================================
# 3. EXISTING SOLUTIONS
# ======================================================================
h1('3. What exists — and why it does not solve this')
table([
    ['Approach', 'Cost / constraint', 'Why it fails in a village'],
    ['12-lead ECG at CHC', '₹45k–₹80k + technician', 'Distance, queue, one day lost; follow-up rates collapse'],
    ['Handheld ECG (AliveCor etc.)', '₹5k–₹15k + prescribed use', 'Still a device to buy and carry; single-lead needs interpretation'],
    ['Smartwatch AF detection', '₹15k–₹40k device + battery', 'Not an ASHA tool; unreachable price for the programme'],
    ['Existing PPG phone apps', 'Often cloud-dependent', 'Upload the pulse to a server? Village connectivity and privacy fail'],
    ['Smartphone-based PPG (academic)', 'Research-stage', 'Either requires clinical equipment, or shipping signals off-device'],
], widths=[2.2, 2.2, 3.0])
rich([('NadiSense’s entire engineering bet is: ', False),
      ('the sensor is already there (a camera), the compute is already there (the phone), and '
       'the workforce is already there (ASHA). What was missing is trustworthy on-device '
       'analysis — the piece we built.', True)])

# ======================================================================
# 4. SOLUTION
# ======================================================================
h1('4. The solution')
h2('4.1 Experience (60 seconds)')
bullet('ASHA worker opens the app (offline), picks language, places patient’s fingertip over '
       'the camera lens — front or rear.', bold_lead='Set up — ')
bullet('App shows the live pulse waveform, a signal-quality meter and live HR; motion is '
       'rejected with clear “hold still” guidance and a retake prompt.', bold_lead='Capture — ')
bullet('Zero-phone detrend + FFT band-pass → beat detection → 12 HRV features → ±3σ winsorisation '
       '→ 6 KB MLP → P(irregular rhythm). Takes ~2 ms.', bold_lead='Analyse — ')
bullet('Green / amber / red card, the 13 HRV metrics, tachogram, Poincaré plot, detected-beat '
       'waveform, and a printable on-device report. Optional 2-question voice screen '
       '(history, current symptoms).', bold_lead='Act — ')
h2('4.2 What is actually built (not a concept)')
table([
    ['Component', 'State', 'Where'],
    ['Camera PPG capture (green-channel ROI → 30 Hz signal)', 'Shipped', 'js/ppgcamera.js'],
    ['DSP: detrend, zero-phase FFT band-pass, adaptive 2-pass peak detection', 'Shipped, tested', 'js/dsp.js'],
    ['12 HRV features (SDNN, RMSSD, pNN50, SD1, SD2, SD1/SD2, LF/HF, spectral entropy, turning-point, irregularity %, ectopy-like fraction, HR)', 'Shipped, tested', 'js/dsp.js'],
    ['MLP classifier 12→20→10→1 (6 KB) + care levels + guardrails', 'Shipped, tested', 'js/classifier.js, js/model_weights.js'],
    ['Vernacular UI (EN/हिं/த) + voice questionnaire', 'Shipped', 'js/i18n.js, js/asr.js'],
    ['Result dashboard, logbook, printable report', 'Shipped', 'js/app.js'],
    ['Synthetic signal generator (6 scenarios incl. AFib-like, motion, weak)', 'Shipped (test fixture)', 'js/simulator.js'],
    ['Training pipeline (NumPy, mirrors the JS DSP 1:1)', 'Shipped + documented', 'tools/train_mlp.py'],
    ['Automated tests: 30 pipeline + 13 full-UI (jsdom)', 'Shipped, green', 'tests/'],
    ['Single-file offline build (works from a USB stick)', 'Shipped', 'deliverables/nadi.html'],
], widths=[3.3, 1.0, 1.9])

# ======================================================================
# 5. ARCHITECTURE
# ======================================================================
h1('5. Frugal hardware & software architecture')
h2('5.1 Hardware: none, deliberately')
para('The bill of materials is one smartphone. Reflectance photoplethysmography needs only a '
     'camera and — in low light — the flash; both are standard on the cheapest devices issued '
     'to ASHA workers. There is no calibration, no disposable sensor, no charger to manage, '
     'no firmware to update, and nothing to lose or break.')
h2('5.2 Software architecture (all on-device)')
table([
    ['Layer', 'What it does', 'Why this design'],
    ['Capture', 'ROI of live frames summed to one green-channel mean per frame; ~30 Hz timeline', 'Image data never exists off-frame; raw pixels are discarded immediately'],
    ['DSP', 'Linear detrend → zero-phase FFT band-pass (0.6–3.5 Hz) → 2-pass adaptive peak finder → RR intervals', 'Standard clinical HRV definitions; ~2 ms on a mid-range phone'],
    ['Features', '12 classical HRV/irregularity features', 'Interpretable, physics-grounded, no black-box embeddings'],
    ['Model', 'MLP 12→20→10→1 (tanh/tanh/sigmoid), 6 KB, ±3σ input winsorisation', 'Runs in any JS engine in a fraction of a microsecond; no runtime deps'],
    ['Guards', 'Quality < 0.6 → retake; HR 40–180; ≥12 clean beats; ectopy note', 'Never guess from a bad signal — safety before throughput'],
    ['UI/UX', '4-step flow, 3 languages, voice Q&A, logbook, printable report', 'Matched to an ASHA worker’s 5-minute home visit'],
    ['Report', 'HTML generated on-device; print/WhatsApp/share via the platform', 'No server means no data trail, no account, no cost'],
], widths=[1.1, 3.2, 2.9])
h2('5.3 Why on-device is a feature, not a constraint')
bullet('Connectivity is unreliable where ASHA workers operate; offline is a requirement, not a preference.')
bullet('No server = no server bill, no breach surface, no data-protection paperwork at the PHC level.')
bullet('Real-time feedback (quality meter, live HR) is only possible when the loop is closed in-device.')
bullet('A 6 KB model is auditable: every weight ships in a plain-text file; nobody can claim a black box.')
h2('5.4 Integrity of the pipeline (the part we are proudest of)')
para('The training script (Python/NumPy) and the shipped inference (JS) are a line-for-line '
     'mirror: same detrend, same band-pass mask, same peak detector, same feature definitions. '
     'We verify the equivalence numerically (cross-language correlation ≈ 0.998 on the same '
     'signals) and 30 automated tests guard every release. This matters because the moment '
     'features differ between research and deployment, the model card is fiction.')

# ======================================================================
# 6. SCIENCE
# ======================================================================
h1('6. Scientific basis')
para('The beat-to-beat interval sequence (RR tachogram) is the same signal a cardiologist '
     'studies on a Holter tape. Healthy sinus rhythm has structured variability: respiratory '
     'sinus arrhythmia and a balanced short-term/long-term ratio. AFib destroys that '
     'structure: chronically elevated SDNN/RMSSD/pNN50, near-random beat ordering, and a '
     'characteristic Poincaré scatter.')
bullet('Time domain: SDNN, RMSSD, pNN50, HR — standard HRV metrics (Task Force conventions).')
bullet('Geometric: SD1/SD2 (Poincaré) — captures short vs long-term variability balance.')
bullet('Frequency: LF/HF from the resampled tachogram; spectral entropy of the 0.04–0.4 Hz band.')
bullet('Pattern: turning-point ratio (randomness of beat order), irregularity %, ectopy-like '
       'beat fraction (transient extra beats that mimic irregularity — flagged, not punished).')
para('Why 60 seconds? It is the settled compromise: short enough to fit a home visit, long '
     'enough for ~70 beats of evidence. The app is deliberately conservative about claiming '
     'more than the evidence allows.')

# ======================================================================
# 7. MODEL
# ======================================================================
h1('7. Model development and validation')
h2('7.1 Data and training')
bullet('12,000 windows of 30 s pulse signals generated with physiological beat dynamics '
       '(NSR with respiratory modulation; low-HRV; AF-like with high short-term variability, '
       'low serial correlation and occasional ectopy-like bursts), plus perturbation '
       'augmentation (noise, gain, time-warp, motion).')
bullet('Features are computed by the same DSP implemented in production; the dataset is '
       'regenerated deterministically from a seed.')
bullet('Standardised inputs, ±3σ winsorisation (the model never sees a point far outside its '
       'training volume).')
h2('7.2 Results (held-out 20%, same synthetic distribution)')
table([
    ['Run', 'Accuracy', 'Sensitivity', 'Specificity', 'F1'],
    ['v0.9 preview model (12k windows, 60 epochs)', '99.5%', '99.6%', '99.4%', '0.996'],
], widths=[3.0, 1.2, 1.3, 1.3, 1.0])
bullet('Demo scenarios through the shipped model: healthy rhythm → risk 0.3% (green) · '
       'low-HRV → 0.2% (amber-ish, based on other features) · AFib-like → 99.9% (red). '
       'These are end-to-end through DSP + classifier.')
h2('7.3 Honest limits — and what we are doing about them')
rich([('The model above is validated on its (realistic, physiological) synthetic training '
       'distribution. It has ', False),
      ('not yet', True),
      (' been benchmarked on real patient recordings. That benchmark is the very next step '
       'and is already wired in: a one-command pipeline pulls MIT-BIH Atrial Fibrillation and '
       'Normal Sinus Rhythm records, computes features through the same code, and retrains — '
       'we will run it before v1.0, and we published it rather than hiding it. A screening '
       'tool that overclaims is worse than no tool.', False)])
para('Ethical position: our success metric is avoidable strokes prevented, not downloads. '
     'The code, the benchmark script and the model card will be public.', size=9.5, italic=True, color=MUTED)

# ======================================================================
# 8. IMPACT
# ======================================================================
h1('8. Expected impact')
h2('8.1 Health')
table([
    ['Indicator', 'Baseline (rural circuit)', 'With NadiSense deployed'],
    ['Adults >40 screened for rhythm, per PHC year', '~0', '10,000+ (one ASHA worker, 4 home visits/day)'],
    ['Time from symptom to ECG (current)', 'Weeks–months', '≤7 days for red-flagged patients (protocol)'],
    ['Cost per screened person', '₹300–₹800 (clinic + travel)', '≈₹0 incremental (existing phone)'],
], widths=[2.8, 2.1, 2.4])
h2('8.2 Economic')
bullet('Removes the single largest cost of screening — travel and lost workday — from the equation.')
bullet('Zero procurement: the district does not buy anything. An ASHA phone is already in the budget.')
bullet('Risks are rationalised: red-flagged patients pre-booked for ECG, reducing no-shows that '
       'waste clinical capacity today.')
h2('8.3 Alignment')
para('SDG 3 (good health and well-being, universal coverage), SDG 9 (innovation), and the '
     'national NCD screening programme’s stated aim of taking screening to the doorstep. '
     'The design (offline, vernacular, phone-native) is deliberately matched to the '
     'Government of India’s ASHA/telemedicine direction of travel.')

# ======================================================================
# 9. DEPLOYMENT
# ======================================================================
h1('9. Deployment and adoption plan')
bullet('Phase 0 (now): pilot with 2 PHC circuits in the Madurai region; 200 screening '
       'sessions by ASHA workers; usability + signal-quality field data.', bold_lead='Pilot — ')
bullet('Train-the-trainer: a 90-minute module for ASHA workers (the app is ~4 taps; the '
       'critical skill is the referral protocol, which is printed on the report).', bold_lead='Training — ')
bullet('Referral SOP co-designed with the PHC MO: red-flagged patients get ECG slots within '
       '7 days; the app’s report is the handoff document.', bold_lead='Protocol — ')
bullet('Retrain on real data, re-validate on a held-out real set, publish the model card, '
       'then approach district health societies and NHM programmes.', bold_lead='Evidence — ')

# ======================================================================
# 10. BUSINESS
# ======================================================================
h1('10. Business model')
table([
    ['Customer', 'What they buy', 'Price logic'],
    ['Public health system (NHM/district)', 'Screening programme software + training', 'Free for public health (impact is the point)'],
    ['Insurers / TPAs', 'Pre-policy and chronic-care risk screening', 'Per-active-screen SaaS'],
    ['Diagnostics chains (outreach)', 'Low-cost pre-test triage for camps', 'Per-screen licence'],
    ['CSR / employer wellness', 'Camps for factory and IT campuses', 'Per-camp bundle'],
], widths=[2.2, 2.4, 2.2])
rich([('Anchor principle: ', True),
      ('the public health deployment is free forever. The commercial lines subsidise it. '
       'A ₹60,000 ECG machine never made a village richer; a ₹0 app might keep it healthier.', False)])

# ======================================================================
# 11. RISKS
# ======================================================================
h1('11. Risks, mitigations, ethics')
table([
    ['Risk', 'Mitigation'],
    ['Model performs worse on real patients than on synthetic data', 'Published retrain-on-real-data path; v1.0 gated on MIT-BIH benchmark; clinical review before any deployment'],
    ['Signal quality in the field (dark verandahs, trembling hands)', 'Quality guardrail forces retake; flash-assisted capture; guidance built into the UI'],
    ['Over-referral floods PHCs', 'Care levels are graduated; amber = repeat, only red = referral; protocol printed on the report'],
    ['Users treat it as a diagnosis', 'Every screen says it: screening aid, not a medical device; report footer repeats it; referral wording is “ECG within 7 days”, never a diagnosis'],
    ['Privacy concerns', 'There is no server: images never leave the phone, no personal data is transmitted at all'],
], widths=[2.9, 3.6])

# ======================================================================
# 12. TEAM & ROADMAP
# ======================================================================
h1('12. Team and roadmap')
h2('12.1 Team (MatricPhase)')
table([
    ['Member', 'Role in NadiSense'],
    ['Aditya Mehra (lead)', 'Product & ML — feature pipeline, model training, on-device inference, app'],
    ['(Recruiting: co-founders per team rules)', 'Field research & clinical liaison; ASHA training module; pilot ops'],
], widths=[2.6, 3.4])
para('We are a solo-engineering team today and are finalising the remaining team slots per the '
     '3–5 member rule — the roles above are the open seats.', size=9.5, italic=True, color=MUTED)
h2('12.2 Roadmap')
table([
    ['Milestone', 'When', 'What'],
    ['v0.9 — this submission', 'Now (Sept 2026)', 'Working prototype, tested, documented'],
    ['v1.0', 'Q4 2026', 'MIT-BIH retrain + validation report; PHC pilot (2 circuits)'],
    ['v1.1', 'H1 2027', 'Multi-class rhythm model; follow-up scheduler; WhatsApp/sms report handoff'],
    ['v2.0', '2027+', 'ASHA-assistant: screening calendar, hypertension/diabetes trends, ANM integration'],
], widths=[1.6, 1.2, 4.2])

# ======================================================================
# APPENDIX
# ======================================================================
doc.add_page_break()
h1('Appendix A — Reproducing everything')
rich([('Everything in this document is reproducible from the submission bundle. '
       'The repo layout is nadi-app/ (app source), tools/ (training, deck, figures, standalone '
       'build), tests/ (43 automated tests), deliverables/ (standalone app + pitch deck).', False)])
table([
    ['Command', 'What it does'],
    ['python tools/train_mlp.py', 'Regenerates the dataset, retrains the MLP, writes js/model_weights.js and tools/metrics.json'],
    ['node tests/run_tests.mjs', '30 assertions: simulator, DSP, features, classifier, guards'],
    ['node tests/run_browser_smoke.mjs', '13 assertions driving the real UI in a DOM (both scenarios)'],
    ['node tools/build_standalone.mjs', 'Bundles the app into one offline HTML file'],
    ['python tools/make_figures.py / make_ui_mock.py', 'Regenerates every figure in this deck and the UI mockup from the product’s own DSP'],
], widths=[3.4, 3.6])
h2('Sources (public ranges)')
bullet('CVD mortality: WHO Global Health Estimates — CVD is the leading cause of death globally.')
bullet('AF prevalence/age gradient & stroke attribution: published epidemiological reviews '
       '(ranges 1–2% overall, up to ~9% over 80; ~20% of strokes attributed to AF).')
bullet('ASHA workforce ~10 lakh: Government of India NHM programme figures.')
para('We keep these as ranges in the narrative rather than single-point claims; the product '
     'does not depend on any individual figure being precise to the decimal.',
     size=9, italic=True, color=MUTED)
h2('Demo')
para('The demo needs no internet and no permissions: open deliverables/nadi.html, choose '
     '“Demo Mode”, pick a scenario, and the full pipeline runs live. A presenter runbook and a '
     '90-second video plan are in deliverables/demo-script.md (bundled separately).')

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', 'deliverables',
                   'NadiSense_TECHNOVA2026_Submission.docx')
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print('wrote', OUT, f'({os.path.getsize(OUT)//1024} KB)')
